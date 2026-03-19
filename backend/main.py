import os
import json
import uuid
import asyncio
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional
from io import BytesIO

from fastapi import FastAPI, HTTPException, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import httpx
from dotenv import load_dotenv

# For PDF & DOCX
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.colors import HexColor
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()

NVIDIA_API_KEY = os.getenv("NVIDIA_API_KEY")
NIM_BASE_URL = "https://integrate.api.nvidia.com/v1"

app = FastAPI(title="ResumeAI Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# In-memory session store (auto-expiring roughly)
# In a real app we'd use a background task to clean this up, but this is stateless-like.
SESSION_STORE: Dict[str, Dict[str, Any]] = {}

# Rate limit tracking: { model_name: timestamp_when_unblocked }
RATE_LIMITS: Dict[str, datetime] = {}

# Fallback Chains
RESEARCHER_CHAIN = [
    "openai/gpt-oss-120b",
    "z-ai/glm5",
    "deepseek-ai/deepseek-r1",
    "moonshotai/kimi-k2.5",
    "nvidia/llama-3.1-nemotron-ultra-253b-v1"
]

WRITER_CHAIN = [
    "openai/gpt-oss-120b",
    "z-ai/glm5",
    "mistralai/mistral-small-4-119b-2603",
    "qwen/qwen3.5-397b-a17b",
    "nvidia/nemotron-3-super-120b-a12b"
]

INTEL_CHAIN = [
    "z-ai/glm5",
    "deepseek-ai/deepseek-r1",
    "openai/gpt-oss-120b",
    "meta/llama-3.1-405b-instruct",
    "nvidia/llama-3.3-nemotron-super-49b-v1"
]

async def call_nim(chain: List[str], system_prompt: str, user_prompt: str, max_retries=2) -> Dict[str, Any]:
    """Helper to call NIM with fallback chain and JSON parsing."""
    if not NVIDIA_API_KEY:
        raise HTTPException(status_code=500, detail="NVIDIA_API_KEY not configured")

    headers = {
        "Authorization": f"Bearer {NVIDIA_API_KEY}",
        "Content-Type": "application/json"
    }

    async with httpx.AsyncClient(timeout=60.0) as client:
        for attempt in range(max_retries):
            # Select model
            selected_model = None
            now = datetime.now()
            
            for model in chain:
                blocked_until = RATE_LIMITS.get(model)
                if not blocked_until or now > blocked_until:
                    selected_model = model
                    break
            
            if not selected_model:
                # All models rate-limited, wait for the soonest one
                soonest_unblocked = min(RATE_LIMITS.values())
                wait_seconds = (soonest_unblocked - now).total_seconds()
                if wait_seconds > 0:
                    await asyncio.sleep(min(wait_seconds, 5.0)) # cap wait time
                selected_model = chain[0] # Try primary anyway after short wait

            payload = {
                "model": selected_model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                "temperature": 0.2,
                "max_tokens": 2048,
                "response_format": {"type": "json_object"}
            }

            try:
                response = await client.post(f"{NIM_BASE_URL}/chat/completions", headers=headers, json=payload)
                if response.status_code == 429:
                    RATE_LIMITS[selected_model] = datetime.now() + timedelta(seconds=61)
                    continue # Try next attempt which will pick next model

                response.raise_for_status()
                data = response.json()
                content = data["choices"][0]["message"]["content"]
                
                # Try parsing JSON
                try:
                    # Strip markdown blocks if present
                    if content.startswith("```json"):
                        content = content[7:-3]
                    elif content.startswith("```"):
                        content = content[3:-3]
                    return json.loads(content)
                except json.JSONDecodeError:
                    if attempt == max_retries - 1:
                        raise ValueError("Failed to parse JSON from AI response")
                    continue # Retry on bad JSON
            except Exception as e:
                print(f"Error with model {selected_model}: {e}")
                if attempt == max_retries - 1:
                    raise HTTPException(status_code=500, detail="AI generation failed after retries")

    raise HTTPException(status_code=500, detail="All fallback models exhausted or failed")


class AnalyzeRequest(BaseModel):
    company: str
    jobTitle: str
    jobDescription: str

@app.post("/api/analyze")
async def analyze_job(req: AnalyzeRequest):
    system_prompt = """You are a professional resume analyst and career coach.
You have access to web search implicit in your vast knowledge base. 
Analyze the company and role thoroughly.
Return ONLY valid JSON matching this exact structure:
{
    "company_meta": {
        "name": "string", "size": "string", "location": "string", "industry": "string"
    },
    "match_analysis": {
        "keywords_must": ["str"],
        "keywords_good": ["str"],
        "keywords_bonus": ["str"],
        "keywords_hit_count": 0
    },
    "company_values": ["str"],
    "ats_score": 90,
    "match_score": 85
}"""
    user_prompt = f"Research {req.company} and this role: {req.jobTitle}\nJob Description:\n{req.jobDescription}\nReturn ONLY the requested JSON."

    data = await call_nim(RESEARCHER_CHAIN, system_prompt, user_prompt)
    return data

class GenerateRequest(BaseModel):
    company: str
    jobTitle: str
    jobDescription: str
    name: str
    experience: str
    email: str
    phone: str
    skills: List[str]
    degree: str
    achievement: str
    linkedin: Optional[str] = ""
    github: Optional[str] = ""
    keywords: List[str]

@app.post("/api/generate")
async def generate_resume(req: GenerateRequest):
    sys_writer = """You are an elite resume writer. Write powerful, ATS-optimized resume content.
Every bullet point must start with a strong action verb, include specific metrics, and use keywords naturally.
Return ONLY valid JSON matching:
{
  "professional_summary": "string",
  "experience": [
    { "company": "string", "title": "string", "duration": "string", "bullets": ["string"] }
  ],
  "skills": ["string"],
  "certifications": ["string"]
}"""
    user_writer = f"Candidate: {req.name}\nApplying to: {req.jobTitle} at {req.company}\nExperience: {req.experience}\nSkills: {req.skills}\nDegree: {req.degree}\nBest Achievement: {req.achievement}\nKeywords to include: {req.keywords}\nJob Desc: {req.jobDescription}\nReturn ONLY JSON."

    sys_intel = """You are a career coach. Return ONLY valid JSON matching:
{
  "culture_points": ["str"],
  "interview_tips": ["str"],
  "salary": { "min": "str", "max": "str", "currency": "str", "notes": ["str"] },
  "interview_questions": [ { "question": "str", "type": "Technical|Behavioural|Situational" } ]
}
Make sure exactly 6 interview questions are provided."""
    user_intel = f"For candidate applying to {req.jobTitle} at {req.company}, return the requested JSON."

    res_writer, res_intel = await asyncio.gather(
        call_nim(WRITER_CHAIN, sys_writer, user_writer),
        call_nim(INTEL_CHAIN, sys_intel, user_intel)
    )

    resume_data = {
        "name": req.name,
        "role": req.jobTitle,
        "email": req.email,
        "phone": req.phone,
        "linkedin": req.linkedin,
        "github": req.github,
        "degree": req.degree,
        "summary": res_writer.get("professional_summary", ""),
        "experience": res_writer.get("experience", []),
        "skills": res_writer.get("skills", req.skills),
        "certifications": res_writer.get("certifications", [])
    }

    session_id = str(uuid.uuid4())
    SESSION_STORE[session_id] = resume_data

    # Return everything needed by frontend
    return {
        "session_id": session_id,
        "resume": resume_data,
        "intel": res_intel,
        "scores": {
            "ats": 95, # In actual scenario, derived from analysis step
            "match": 90
        }
    }

class DownloadRequest(BaseModel):
    resume: Dict[str, Any]

@app.post("/api/download/pdf")
async def download_pdf(req: DownloadRequest):
    data = req.resume
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    # Colors
    ink = HexColor("#0a0a0f")
    gold = HexColor("#c9a84c")
    cream = HexColor("#f5f0e8")
    
    # Header block
    c.setFillColor(ink)
    c.rect(0, height - 1.5*inch, width, 1.5*inch, fill=1, stroke=0)
    
    c.setFillColor(cream)
    c.setFont("Helvetica-Bold", 24)
    c.drawString(0.5*inch, height - 0.7*inch, data.get("name", "Name"))
    
    c.setFillColor(gold)
    c.setFont("Helvetica", 10)
    c.drawString(0.5*inch, height - 0.9*inch, data.get("role", "Role").upper())
    
    # divider
    c.setStrokeColor(gold)
    c.line(0.5*inch, height - 1.05*inch, width - 0.5*inch, height - 1.05*inch)
    
    # contact
    c.setFillColor(cream)
    c.setFont("Helvetica", 8)
    contact_parts = []
    if data.get("email"): contact_parts.append(data["email"])
    if data.get("phone"): contact_parts.append(data["phone"])
    if data.get("linkedin"): contact_parts.append(data["linkedin"])
    c.drawString(0.5*inch, height - 1.3*inch, " | ".join(contact_parts))
    
    # Body
    text_y = height - 2*inch
    c.setFillColor(ink)
    
    # Summary
    c.setFont("Helvetica-Bold", 10)
    c.setFillColor(gold)
    c.drawString(0.5*inch, text_y, "PROFESSIONAL SUMMARY")
    c.line(0.5*inch, text_y - 2, 2.5*inch, text_y - 2)
    
    text_y -= 0.25*inch
    c.setFillColor(ink)
    c.setFont("Helvetica", 9)
    # Simple text split (reportlab canvas is basic)
    # A real world app would use platypus for wrapping, but simple string slicing for demo:
    summary = data.get("summary", "")
    for i in range(0, len(summary), 90):
        c.drawString(0.5*inch, text_y, summary[i:i+90])
        text_y -= 12
    
    text_y -= 0.2*inch
    
    # Experience
    c.setFont("Helvetica-Bold", 10)
    c.setFillColor(gold)
    c.drawString(0.5*inch, text_y, "EXPERIENCE")
    c.line(0.5*inch, text_y - 2, 1.5*inch, text_y - 2)
    text_y -= 0.25*inch
    
    for job in data.get("experience", []):
        c.setFillColor(ink)
        c.setFont("Helvetica-Bold", 9)
        c.drawString(0.5*inch, text_y, job.get("company", ""))
        c.setFont("Helvetica", 8)
        c.drawRightString(width - 0.5*inch, text_y, job.get("duration", ""))
        text_y -= 12
        
        c.setFillColor(gold)
        c.setFont("Helvetica-Oblique", 9)
        c.drawString(0.5*inch, text_y, job.get("title", ""))
        text_y -= 14
        
        c.setFillColor(ink)
        c.setFont("Helvetica", 9)
        for bullet in job.get("bullets", []):
            if text_y < 1*inch:
                c.showPage()
                text_y = height - 1*inch
                
            bullet_lines = [bullet[i:i+100] for i in range(0, len(bullet), 100)]
            c.drawString(0.6*inch, text_y, f"> {bullet_lines[0]}")
            text_y -= 12
            for bline in bullet_lines[1:]:
                c.drawString(0.7*inch, text_y, bline)
                text_y -= 12
        text_y -= 0.1*inch
        
    c.showPage()
    c.save()
    
    buffer.seek(0)
    headers = {"Content-Disposition": f'attachment; filename="{data.get("name", "Resume").replace(" ", "_").replace("/", "-")}.pdf"'}
    return StreamingResponse(buffer, media_type="application/pdf", headers=headers)


@app.post("/api/download/word")
async def download_word(req: DownloadRequest):
    data = req.resume
    doc = Document()
    
    # Header
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(data.get("name", "Name"))
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    role_para = doc.add_paragraph()
    role_run = role_para.add_run(data.get("role", "Role").upper())
    role_run.font.size = Pt(11)
    role_run.font.color.rgb = RGBColor(201, 168, 76) # Gold
    role_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact_parts = []
    if data.get("email"): contact_parts.append(data["email"])
    if data.get("phone"): contact_parts.append(data["phone"])
    if data.get("linkedin"): contact_parts.append(data["linkedin"])
    if data.get("github"): contact_parts.append(data["github"])
    
    contact_para = doc.add_paragraph(" | ".join(contact_parts))
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Body
    doc.add_heading("PROFESSIONAL SUMMARY", level=2)
    doc.add_paragraph(data.get("summary", ""))
    
    doc.add_heading("EXPERIENCE", level=2)
    for job in data.get("experience", []):
        p = doc.add_paragraph()
        p.add_run(f"{job.get('company', '')} ").bold = True
        p.add_run(f" | {job.get('duration', '')}")
        
        p_title = doc.add_paragraph()
        run_title = p_title.add_run(job.get("title", ""))
        run_title.italic = True
        run_title.font.color.rgb = RGBColor(201, 168, 76)
        
        for bullet in job.get("bullets", []):
            doc.add_paragraph(bullet, style='List Bullet')
            
    doc.add_heading("TECHNICAL SKILLS", level=2)
    doc.add_paragraph(", ".join(data.get("skills", [])))
    
    doc.add_heading("EDUCATION & CERTIFICATIONS", level=2)
    doc.add_paragraph(data.get("degree", ""))
    for cert in data.get("certifications", []):
        doc.add_paragraph(cert, style='List Bullet')
        
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    headers = {"Content-Disposition": f'attachment; filename="{data.get("name", "Resume").replace(" ", "_").replace("/", "-")}.docx"'}
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)
